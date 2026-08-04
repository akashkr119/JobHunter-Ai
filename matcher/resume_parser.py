"""Resume parsing and skill-extraction utilities."""

import re
from pathlib import Path


class ResumeParser:
    """Extract resume text and identify common technical skills."""

    DEFAULT_SKILLS = (
        "python", "java", "javascript", "typescript", "c++", "c#", "sql",
        "html", "css", "selenium", "pytest", "playwright", "robot framework",
        "appium", "jira", "git", "github", "jenkins", "docker", "kubernetes",
        "aws", "azure", "gcp", "flask", "django", "fastapi", "pandas", "numpy",
        "rest api", "api testing", "automation testing", "manual testing", "can",
        "canoe", "canalyzer", "capl", "uds", "automotive",
    )

    SUPPORTED_FORMATS = {".txt", ".md", ".pdf", ".docx"}

    def extract_text(self, resume_path: str | Path) -> str:
        """Extract text from TXT, Markdown, PDF or DOCX resumes."""
        path = Path(resume_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {path}")
        if not path.is_file():
            raise ValueError(f"Resume path is not a file: {path}")

        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            supported = ", ".join(sorted(self.SUPPORTED_FORMATS))
            raise ValueError(
                f"Unsupported resume format: {suffix or 'unknown'}. "
                f"Currently supported: {supported}"
            )

        if suffix in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".pdf":
            text = self._extract_pdf(path)
        else:
            text = self._extract_docx(path)

        return self._clean_text(text)

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        """Extract text from a text-based PDF resume."""
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF parsing requires the pypdf package") from exc

        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(pages)

    @staticmethod
    def _extract_docx(path: Path) -> str:
        """Extract paragraphs and table text from a DOCX resume."""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX parsing requires the python-docx package") from exc

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)

    def extract_skills(
        self,
        text: str,
        skills: list[str] | tuple[str, ...] | None = None,
    ) -> list[str]:
        """Return known skills found in resume text, preserving skill order."""
        normalized_text = self._normalize(text)
        candidates = skills or self.DEFAULT_SKILLS
        found: list[str] = []

        for skill in candidates:
            normalized_skill = self._normalize(str(skill))
            if not normalized_skill:
                continue
            pattern = rf"(?<![a-z0-9]){re.escape(normalized_skill)}(?![a-z0-9])"
            if re.search(pattern, normalized_text) and normalized_skill not in found:
                found.append(normalized_skill)

        return found

    def parse(self, resume_path: str | Path) -> dict:
        """Parse a resume into normalized text and detected skills."""
        path = Path(resume_path)
        text = self.extract_text(path)
        return {
            "path": str(path),
            "format": path.suffix.lower().lstrip("."),
            "text": text,
            "skills": self.extract_skills(text),
        }

    @staticmethod
    def _clean_text(value: str) -> str:
        lines = [re.sub(r"\s+", " ", line).strip() for line in str(value or "").splitlines()]
        return "\n".join(line for line in lines if line).strip()

    @staticmethod
    def _normalize(value: str) -> str:
        """Normalize text for deterministic skill matching."""
        return re.sub(r"\s+", " ", str(value or "").lower()).strip()
