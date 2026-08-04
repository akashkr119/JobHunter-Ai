"""Unit tests for resume parser."""

from unittest.mock import MagicMock, patch

import pytest

from matcher.resume_parser import ResumeParser


def test_resume_parser_instance():
    assert ResumeParser() is not None


def test_has_parse_method():
    assert hasattr(ResumeParser(), "parse")


def test_parse_method_is_callable():
    assert callable(ResumeParser().parse)


def test_extract_text_reads_plain_text_resume(tmp_path):
    resume = tmp_path / "resume.txt"
    resume.write_text("Python Selenium Pytest", encoding="utf-8")
    assert ResumeParser().extract_text(resume) == "Python Selenium Pytest"


def test_extract_skills_is_case_insensitive():
    skills = ResumeParser().extract_skills("PYTHON, Selenium, pytest and Docker")
    assert {"python", "selenium", "pytest", "docker"}.issubset(skills)


def test_extract_skills_detects_automotive_skills():
    skills = ResumeParser().extract_skills(
        "Automotive testing experience with CAN, CANoe, CAPL and UDS."
    )
    assert {"automotive", "can", "canoe", "capl", "uds"}.issubset(skills)


def test_parse_returns_structured_resume(tmp_path):
    resume = tmp_path / "resume.txt"
    resume.write_text(
        "QA Automation Engineer with Python, Selenium, Pytest and Jenkins.",
        encoding="utf-8",
    )
    result = ResumeParser().parse(resume)
    assert result["path"] == str(resume)
    assert result["format"] == "txt"
    assert "QA Automation Engineer" in result["text"]
    assert {"python", "selenium", "pytest", "jenkins"}.issubset(result["skills"])


def test_missing_resume_raises_file_not_found(tmp_path):
    with pytest.raises(FileNotFoundError):
        ResumeParser().parse(tmp_path / "missing.txt")


def test_unsupported_resume_format_rejected(tmp_path):
    resume = tmp_path / "resume.rtf"
    resume.write_text("Python", encoding="utf-8")
    with pytest.raises(ValueError, match="Unsupported resume format"):
        ResumeParser().parse(resume)


def test_markdown_resume_is_supported(tmp_path):
    resume = tmp_path / "resume.md"
    resume.write_text("# Skills\nPython\nSelenium", encoding="utf-8")
    result = ResumeParser().parse(resume)
    assert result["format"] == "md"
    assert {"python", "selenium"}.issubset(result["skills"])


def test_docx_resume_parsing(tmp_path):
    from docx import Document

    resume = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("QA Engineer with Python Selenium Pytest")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Docker Jenkins"
    document.save(resume)

    result = ResumeParser().parse(resume)
    assert result["format"] == "docx"
    assert "QA Engineer" in result["text"]
    assert {"python", "selenium", "pytest", "docker", "jenkins"}.issubset(
        result["skills"]
    )


def test_pdf_resume_parsing(tmp_path):
    resume = tmp_path / "resume.pdf"
    resume.write_bytes(b"placeholder")

    page = MagicMock()
    page.extract_text.return_value = "Python Selenium Pytest Docker"
    reader = MagicMock()
    reader.pages = [page]

    with patch("pypdf.PdfReader", return_value=reader) as pdf_reader:
        result = ResumeParser().parse(resume)

    pdf_reader.assert_called_once_with(str(resume))
    assert result["format"] == "pdf"
    assert {"python", "selenium", "pytest", "docker"}.issubset(result["skills"])


def test_text_cleanup_removes_blank_lines_and_extra_spaces(tmp_path):
    resume = tmp_path / "resume.txt"
    resume.write_text("Python   Selenium\n\n   Pytest  ", encoding="utf-8")
    text = ResumeParser().extract_text(resume)
    assert text == "Python Selenium\nPytest"
